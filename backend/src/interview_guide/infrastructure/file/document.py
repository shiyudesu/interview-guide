from __future__ import annotations

import asyncio
import logging
import shutil
import subprocess
import tempfile
from io import BytesIO
from pathlib import Path
from typing import Protocol
from urllib.parse import quote

from docx import Document
from pdfminer.high_level import extract_text
from pdfminer.layout import LAParams
from pypdf import PdfReader

from interview_guide.common.api.models import java_utf16_length
from interview_guide.common.config.settings import Settings
from interview_guide.common.errors import BusinessException, ErrorCode
from interview_guide.common.runtime import BlockingExecutor
from interview_guide.infrastructure.file.text import clean_text
from interview_guide.infrastructure.file.validation import DOCUMENT_MAX_UTF16_UNITS

logger = logging.getLogger(__name__)


class OfficeConverter(Protocol):
    def convert_to_docx(self, data: bytes, suffix: str) -> bytes: ...


class LibreOfficeConverter:
    def __init__(
        self,
        timeout_seconds: float,
        max_output_bytes: int,
        executable: str | None = None,
    ) -> None:
        self._timeout_seconds = timeout_seconds
        self._max_output_bytes = max_output_bytes
        self._executable = executable or shutil.which("libreoffice") or shutil.which("soffice")

    def convert_to_docx(self, data: bytes, suffix: str) -> bytes:
        if self._executable is None:
            raise RuntimeError("LibreOffice executable is not installed")
        with tempfile.TemporaryDirectory(prefix="interview-guide-document-") as directory:
            workdir = Path(directory)
            source = workdir / f"source{suffix}"
            output = workdir / "source.docx"
            profile = workdir / "profile"
            source.write_bytes(data)
            result = subprocess.run(
                [
                    self._executable,
                    "--headless",
                    f"-env:UserInstallation=file://{quote(str(profile))}",
                    "--convert-to",
                    "docx",
                    "--outdir",
                    str(workdir),
                    str(source),
                ],
                capture_output=True,
                check=False,
                timeout=self._timeout_seconds,
            )
            if result.returncode != 0 or not output.exists():
                detail = result.stderr.decode("utf-8", errors="replace").strip()
                raise RuntimeError(f"LibreOffice conversion failed: {detail}")
            if output.stat().st_size > self._max_output_bytes:
                raise RuntimeError("LibreOffice conversion output exceeds limit")
            return output.read_bytes()


def decode_text(data: bytes) -> str:
    failures: list[UnicodeDecodeError] = []
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "big5"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError as error:
            failures.append(error)
    try:
        return data.decode("latin-1")
    except UnicodeDecodeError as error:
        failures.append(error)
    raise failures[-1]


def extract_docx(data: bytes) -> str:
    document = Document(BytesIO(data))
    blocks: list[str] = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            blocks.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(blocks)


def extract_pdf(data: bytes) -> str:
    reader = PdfReader(BytesIO(data))
    if reader.is_encrypted:
        raise RuntimeError("PDF 文件已加密")
    _ = len(reader.pages)
    return extract_text(
        BytesIO(data),
        laparams=LAParams(boxes_flow=0.5, all_texts=False),
    )


class DocumentParser:
    def __init__(self, office_converter: OfficeConverter) -> None:
        self._office_converter = office_converter

    def parse(
        self,
        data: bytes,
        filename: str | None,
        content_type: str | None = None,
    ) -> str:
        if not data:
            return ""
        try:
            raw_text = self._extract(data, filename, content_type)
            if java_utf16_length(raw_text) > DOCUMENT_MAX_UTF16_UNITS:
                raise RuntimeError("文档内容超过最大字符限制")
            return clean_text(raw_text)
        except BusinessException:
            raise
        except Exception as error:
            logger.exception("document parsing failed filename=%s", filename)
            raise BusinessException(
                ErrorCode.INTERNAL_ERROR,
                f"文件解析失败: {error}",
            ) from error

    def _extract(
        self,
        data: bytes,
        filename: str | None,
        content_type: str | None,
    ) -> str:
        suffix = Path(filename or "").suffix.lower()
        normalized_type = (content_type or "").lower()
        if suffix == ".pdf" or "pdf" in normalized_type:
            return extract_pdf(data)
        if suffix == ".docx" or "wordprocessingml" in normalized_type:
            return extract_docx(data)
        if suffix in {".doc", ".rtf"} or normalized_type in {
            "application/msword",
            "application/rtf",
        }:
            converted = self._office_converter.convert_to_docx(
                data,
                suffix or ".doc",
            )
            return extract_docx(converted)
        return decode_text(data)


class AsyncDocumentParser:
    def __init__(
        self,
        parser: DocumentParser,
        executor: BlockingExecutor,
        concurrency: int,
    ) -> None:
        self._parser = parser
        self._executor = executor
        self._semaphore = asyncio.Semaphore(concurrency)

    async def parse(
        self,
        data: bytes,
        filename: str | None,
        content_type: str | None = None,
    ) -> str:
        async with self._semaphore:
            return await self._executor.run(
                self._parser.parse,
                data,
                filename,
                content_type,
            )


def create_document_parser(
    settings: Settings,
    executor: BlockingExecutor,
) -> AsyncDocumentParser:
    converter = LibreOfficeConverter(
        settings.document_conversion_timeout_seconds,
        settings.document_conversion_max_bytes,
    )
    return AsyncDocumentParser(
        DocumentParser(converter),
        executor,
        settings.document_worker_count,
    )

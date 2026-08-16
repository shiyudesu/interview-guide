from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document
from pypdf import PdfWriter
from reportlab.pdfgen import canvas

from interview_guide.common.errors import BusinessException
from interview_guide.common.runtime import BlockingExecutor
from interview_guide.infrastructure.file.document import (
    AsyncDocumentParser,
    DocumentParser,
    decode_text,
)


def docx_bytes(text: str) -> bytes:
    document = Document()
    document.add_paragraph(text)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def pdf_bytes(text: str) -> bytes:
    output = BytesIO()
    document = canvas.Canvas(output)
    document.drawString(72, 720, text)
    document.save()
    return output.getvalue()


class FakeOfficeConverter:
    def __init__(self, converted: bytes) -> None:
        self.converted = converted
        self.calls: list[tuple[bytes, str]] = []

    def convert_to_docx(self, data: bytes, suffix: str) -> bytes:
        self.calls.append((data, suffix))
        return self.converted


def test_txt_markdown_and_explicit_encoding_order() -> None:
    parser = DocumentParser(FakeOfficeConverter(b""))

    assert parser.parse("姓名：张三".encode(), "resume.txt", "text/plain") == "姓名：张三"
    assert "# 标题" in parser.parse(
        "# 标题".encode(),
        "knowledge.md",
        "text/markdown",
    )
    assert decode_text("中文".encode("gb18030")) == "中文"


def test_docx_extracts_paragraphs_and_tables() -> None:
    document = Document()
    document.add_paragraph("个人简历")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "姓名"
    table.cell(0, 1).text = "李四"
    output = BytesIO()
    document.save(output)
    parser = DocumentParser(FakeOfficeConverter(b""))

    result = parser.parse(
        output.getvalue(),
        "resume.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    assert "个人简历" in result
    assert "姓名\t李四" in result


def test_pdf_uses_pdfminer_and_rejects_encryption() -> None:
    parser = DocumentParser(FakeOfficeConverter(b""))
    assert "Resume PDF" in parser.parse(
        pdf_bytes("Resume PDF"),
        "resume.pdf",
        "application/pdf",
    )

    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    writer.encrypt("secret")
    encrypted = BytesIO()
    writer.write(encrypted)

    with pytest.raises(BusinessException, match="PDF 文件已加密"):
        parser.parse(
            encrypted.getvalue(),
            "encrypted.pdf",
            "application/pdf",
        )


@pytest.mark.parametrize(
    ("filename", "content_type"),
    [
        ("legacy.doc", "application/msword"),
        ("legacy.rtf", "application/rtf"),
    ],
)
def test_doc_and_rtf_use_controlled_office_conversion(
    filename: str,
    content_type: str,
) -> None:
    converter = FakeOfficeConverter(docx_bytes("转换后的正文"))
    parser = DocumentParser(converter)

    assert parser.parse(b"legacy", filename, content_type) == "转换后的正文"
    assert converter.calls == [(b"legacy", f".{filename.rsplit('.', 1)[1]}")]


def test_character_limit_is_checked_before_cleaning() -> None:
    parser = DocumentParser(FakeOfficeConverter(b""))
    content = ("x" * (5 * 1024 * 1024 + 1)).encode()

    with pytest.raises(BusinessException, match="文档内容超过最大字符限制"):
        parser.parse(content, "oversized.txt", "text/plain")


@pytest.mark.asyncio
async def test_async_parser_uses_bounded_blocking_executor() -> None:
    executor = BlockingExecutor(max_workers=1)
    parser = AsyncDocumentParser(
        DocumentParser(FakeOfficeConverter(b"")),
        executor,
        concurrency=1,
    )

    assert await parser.parse(b"fixed", "sample.txt", "text/plain") == "fixed"
    await executor.shutdown()
